import os
from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

class KnowledgeBase:
    def __init__(self, persist_directory="./chroma_db", model_name="nomic-embed-text"):
        self.persist_directory = persist_directory
        self.model_name = model_name
        self.embeddings = None
        self.vectorstore = None
    
    def _init_embeddings(self):
        if self.embeddings is None:
            from langchain_ollama import OllamaEmbeddings
            self.embeddings = OllamaEmbeddings(model=self.model_name)
    
    def _load_or_create_vectorstore(self):
        if self.vectorstore is None:
            from langchain_chroma import Chroma
            self._init_embeddings()
            if os.path.exists(self.persist_directory):
                print("加载已有的向量数据库...")
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
            else:
                print("创建新的向量数据库...")
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
    
    def read_pdf(self, file_path):
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    
    def read_docx(self, file_path):
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    
    def read_txt(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def read_document(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self.read_pdf(file_path)
        elif ext == ".docx":
            return self.read_docx(file_path)
        elif ext == ".txt":
            return self.read_txt(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    def split_text(self, text, chunk_size=1000, chunk_overlap=200):
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
        return text_splitter.split_text(text)
    
    def add_documents_from_folder(self, folder_path):
        self._load_or_create_vectorstore()
        documents = []
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            if os.path.isfile(file_path):
                ext = os.path.splitext(filename)[1].lower()
                if ext in [".pdf", ".docx", ".txt"]:
                    try:
                        print(f"正在处理: {filename}")
                        text = self.read_document(file_path)
                        chunks = self.split_text(text)
                        for i, chunk in enumerate(chunks):
                            doc = LangChainDocument(
                                page_content=chunk,
                                metadata={"source": filename, "chunk": i}
                            )
                            documents.append(doc)
                    except Exception as e:
                        print(f"处理 {filename} 时出错: {str(e)}")
        
        if documents:
            self.vectorstore.add_documents(documents)
            self.vectorstore.persist()
            print(f"成功添加 {len(documents)} 个文本块到知识库")
        return documents
    
    def add_document(self, file_path):
        self._load_or_create_vectorstore()
        documents = []
        try:
            filename = os.path.basename(file_path)
            print(f"正在处理: {filename}")
            text = self.read_document(file_path)
            chunks = self.split_text(text)
            for i, chunk in enumerate(chunks):
                doc = LangChainDocument(
                    page_content=chunk,
                    metadata={"source": filename, "chunk": i}
                )
                documents.append(doc)
            
            if documents:
                self.vectorstore.add_documents(documents)
                self.vectorstore.persist()
                print(f"成功添加 {len(documents)} 个文本块到知识库")
        except Exception as e:
            print(f"处理文件时出错: {str(e)}")
        return documents
    
    def retrieve(self, query, k=3):
        self._load_or_create_vectorstore()
        if self.vectorstore is None:
            return []
        results = self.vectorstore.similarity_search(query, k=k)
        return results
    
    def get_vectorstore(self):
        self._load_or_create_vectorstore()
        return self.vectorstore
    
    def get_document_count(self):
        try:
            self._load_or_create_vectorstore()
            if self.vectorstore is None:
                return 0
            return len(self.vectorstore.get()["ids"])
        except:
            return 0
    
    def get_unique_sources(self):
        try:
            self._load_or_create_vectorstore()
            if self.vectorstore is None:
                return []
            metadata = self.vectorstore.get()["metadatas"]
            sources = list(set([doc["source"] for doc in metadata]))
            return sources
        except:
            return []

if __name__ == "__main__":
    kb = KnowledgeBase()
    
    docs_folder = "./documents"
    if os.path.exists(docs_folder):
        kb.add_documents_from_folder(docs_folder)
    else:
        print(f"文件夹 {docs_folder} 不存在，请先创建该文件夹并放入文档")
    
    print("\n测试检索功能：")
    query = "什么是自然语言处理？"
    results = kb.retrieve(query)
    print(f"查询: {query}")
    print("检索结果:")
    for i, doc in enumerate(results):
        print(f"\n结果 {i+1} (来源: {doc.metadata['source']}):")
        print(doc.page_content[:300], "...")