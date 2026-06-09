import streamlit as st
import os
import tempfile
import shutil

st.set_page_config(
    page_title="RAG 本地知识库问答系统",
    page_icon="🤖",
    layout="wide"
)

st.title("🤖 RAG 本地知识库问答系统")

def init_session_state():
    if "documents" not in st.session_state:
        st.session_state.documents = []
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "vectorstore" not in st.session_state:
        st.session_state.vectorstore = None
    if "qa_chain" not in st.session_state:
        st.session_state.qa_chain = None

def load_documents_from_folder(folder_path):
    docs = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path):
            ext = os.path.splitext(filename)[1].lower()
            if ext in [".pdf", ".docx", ".txt"]:
                try:
                    text = read_document(file_path)
                    docs.append({"name": filename, "content": text})
                except:
                    pass
    return docs

def read_pdf(file_path):
    from PyPDF2 import PdfReader
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def read_docx(file_path):
    from docx import Document
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def read_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def read_document(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return read_pdf(file_path)
    elif ext == ".docx":
        return read_docx(file_path)
    elif ext == ".txt":
        return read_txt(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")

def split_text(text, chunk_size=1000, chunk_overlap=200):
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )
    return text_splitter.split_text(text)

def create_vectorstore(documents):
    from langchain_chroma import Chroma
    from langchain_ollama import OllamaEmbeddings
    from langchain.schema import Document as LangChainDocument
    
    chroma_path = "./chroma_db"
    if os.path.exists(chroma_path):
        shutil.rmtree(chroma_path)
    
    embeddings = OllamaEmbeddings(model="nomic-embed-text")
    db = Chroma(persist_directory=chroma_path, embedding_function=embeddings)
    
    all_docs = []
    for doc in documents:
        chunks = split_text(doc["content"])
        for i, chunk in enumerate(chunks):
            lc_doc = LangChainDocument(
                page_content=chunk,
                metadata={"source": doc["name"], "chunk": i}
            )
            all_docs.append(lc_doc)
    
    db.add_documents(all_docs)
    db.persist()
    return db

def create_qa_chain(vectorstore):
    from langchain.chains import ConversationalRetrievalChain
    from langchain_ollama import OllamaLLM
    from langchain.prompts import PromptTemplate
    
    prompt_template = """基于以下参考文档回答问题：

参考文档：
{context}

用户问题：{question}

如果文档中有相关信息，请基于文档内容回答。如果没有相关信息，请明确说"文档中未找到相关答案"。

回答："""

    PROMPT = PromptTemplate(
        template=prompt_template,
        input_variables=["context", "question"]
    )

    retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
    llm = OllamaLLM(model="qwen2:0.5b")

    qa_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=retriever,
        return_source_documents=True,
        combine_docs_chain_kwargs={"prompt": PROMPT},
        verbose=False
    )
    
    return qa_chain

init_session_state()

sidebar = st.sidebar
with sidebar:
    st.header("📚 知识库管理")
    
    uploaded_files = st.file_uploader(
        "上传文档 (PDF/DOCX/TXT)",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        help="支持 PDF、DOCX 和 TXT 格式"
    )
    
    if uploaded_files:
        st.session_state.uploaded_files = uploaded_files
        st.success(f"已选择 {len(uploaded_files)} 个文件")
    
    use_sample_docs = st.checkbox("使用示例文档", value=True)
    
    if st.button("🔨 构建知识库", type="primary"):
        with st.spinner("正在处理文档..."):
            documents = []
            
            if use_sample_docs:
                docs_folder = "./documents"
                if os.path.exists(docs_folder):
                    sample_docs = load_documents_from_folder(docs_folder)
                    documents.extend(sample_docs)
                    st.info(f"已加载 {len(sample_docs)} 个示例文档")
            
            if uploaded_files:
                for uploaded_file in uploaded_files:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
                        tmp_file.write(uploaded_file.getvalue())
                        tmp_file_path = tmp_file.name
                    
                    try:
                        text = read_document(tmp_file_path)
                        documents.append({"name": uploaded_file.name, "content": text})
                        st.success(f"已处理: {uploaded_file.name}")
                    finally:
                        os.unlink(tmp_file_path)
            
            if documents:
                try:
                    st.session_state.vectorstore = create_vectorstore(documents)
                    st.session_state.qa_chain = None
                    st.success(f"✅ 知识库构建完成！共 {len(documents)} 个文档")
                except Exception as e:
                    st.error(f"构建知识库失败: {str(e)}")
            else:
                st.warning("没有文档需要处理")
    
    st.divider()
    
    st.subheader("📊 知识库状态")
    if st.session_state.vectorstore:
        try:
            doc_count = len(st.session_state.vectorstore.get()["ids"])
            st.metric("文本块数量", doc_count)
        except:
            st.info("知识库已构建")
    else:
        st.info("知识库尚未构建")
    
    if st.button("🗑️ 清空对话历史"):
        st.session_state.chat_history = []
        st.success("对话历史已清空")

st.subheader("💬 问答交互")

for i, (question, answer) in enumerate(st.session_state.chat_history):
    with st.chat_message("user"):
        st.write(question)
    with st.chat_message("assistant"):
        st.write(answer)

if question := st.chat_input("请输入您的问题..."):
    with st.chat_message("user"):
        st.write(question)
    
    with st.chat_message("assistant"):
        if not st.session_state.vectorstore:
            st.warning("⚠️ 知识库尚未构建，请先上传文档并点击'构建知识库'")
        else:
            try:
                if st.session_state.qa_chain is None:
                    with st.spinner("初始化问答系统..."):
                        st.session_state.qa_chain = create_qa_chain(st.session_state.vectorstore)
                
                with st.spinner("正在思考..."):
                    result = st.session_state.qa_chain({
                        "question": question,
                        "chat_history": []
                    })
                    
                    answer = result["answer"]
                    source_docs = result["source_documents"]
                    
                    st.write(answer)
                    
                    if source_docs:
                        with st.expander("📖 查看参考来源"):
                            for i, doc in enumerate(source_docs):
                                st.write(f"**来源**: {doc.metadata['source']}")
                                st.write(doc.page_content[:200] + "...")
                                st.divider()
            
            except Exception as e:
                st.error(f"回答时出错: {str(e)}")
                answer = "抱歉，回答过程中出现错误"
    
    st.session_state.chat_history.append((question, answer))